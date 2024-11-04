import os, re
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def format_answer_text(paragraph, answer_text):
    paragraphs = answer_text.split("\n\n")
    
    for para in paragraphs:
        lines = para.splitlines()
        
        for line in lines:
            # Create a new paragraph for each line
            current_paragraph = paragraph.add_paragraph()
            
            # Handle headers starting with ###
            if line.startswith("### "):
                header_run = current_paragraph.add_run(line[4:])
                header_run.bold = True
                header_run.font.size = Pt(14)
                current_paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            # Primary bullet points
            elif line.startswith("1. ") or line.startswith("- "):
                current_paragraph.paragraph_format.left_indent = Pt(20)
                process_text_with_bold(current_paragraph, line)

            # Secondary bullet points with extra indentation
            elif line.startswith("  - "):
                current_paragraph.paragraph_format.left_indent = Pt(40)
                process_text_with_bold(current_paragraph, line)
            
            # Regular text
            else:
                process_text_with_bold(current_paragraph, line)


def process_text_with_bold(paragraph, text):
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            bold_run = paragraph.add_run(part[2:-2])  # Remove the ** markers
            bold_run.bold = True
            bold_run.font.size = Pt(11)
        else:
            normal_run = paragraph.add_run(part)
            normal_run.font.size = Pt(11)