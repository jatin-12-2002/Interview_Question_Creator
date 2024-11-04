from fastapi import FastAPI, Form, Request, File, HTTPException, UploadFile
from fastapi.responses import JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.encoders import jsonable_encoder
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os,re
import aiofiles
from src.helper import llm_pipeline
from fastapi.middleware.cors import CORSMiddleware
from celery import Celery

# Configure Celery
celery = Celery(
    "tasks",
    broker="redis://localhost:6379/0",
    backend="redis://localhost:6379/1"
)
celery.conf.update(
    task_serializer="json",
    result_serializer="json",
    accept_content=["json"],
)

app = FastAPI()

# Allow all origins for testing; specify allowed origins in production
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Serve static files (for styles, docs, etc.)
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")


@app.get("/")
async def index(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


@app.post("/analyze")
async def analyze_pdf(pdf_file: UploadFile = File(...), num_questions: int = Form(...)):
    base_folder = 'static/docs/'
    os.makedirs(base_folder, exist_ok=True)
    pdf_filename = os.path.join(base_folder, pdf_file.filename)

    # Save the uploaded PDF
    try:
        async with aiofiles.open(pdf_filename, 'wb') as out_file:
            content = await pdf_file.read()
            await out_file.write(content)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save PDF: {str(e)}")

    # Enqueue the task with Celery
    task = process_pdf_task.delay(pdf_filename, num_questions)
    return JSONResponse(content=jsonable_encoder({"task_id": task.id}))


@celery.task(bind=True)
def process_pdf_task(self, file_path, num_questions):
    try:
        output_file_name = get_docx(file_path, num_questions)
        return {"status": "completed", "output_file": output_file_name}
    except Exception as e:
        return {"status": "failed", "error": str(e)}


def get_docx(file_path, num_questions):
    pipeline_output = llm_pipeline(file_path, num_questions)
    ques_list = pipeline_output["questions"]
    answer_generation_chain = pipeline_output["answer_generation_chain"]

    base_folder = 'static/output/'
    os.makedirs(base_folder, exist_ok=True)

    output_file_name = "QA.docx"
    output_file = os.path.join(base_folder, output_file_name)

    # Create a DOCX document
    doc = Document()
    doc.add_heading('Interview Questions and Answers', 0)

    for i, question in enumerate(ques_list[:num_questions], start=1):
        # Add Question heading
        question_heading = doc.add_paragraph()
        question_heading.add_run(f"Question {i}:").bold = True
        question_heading.style = 'Heading 2'  # Style to make it look like a header

        # Add Question text
        question_paragraph = doc.add_paragraph(question)
        question_paragraph.style = doc.styles['Normal']

        # Retrieve answer
        try:
            answer = answer_generation_chain.run(question)
        except Exception as e:
            answer = f"Answer generation failed: {str(e)}"

        # Add Answer heading
        answer_heading = doc.add_paragraph()
        answer_heading.add_run("Answer:").bold = True
        answer_heading.style = 'Heading 2'

        # Format Answer text into Markdown-like style
        # answer_paragraph = doc.add_paragraph()
        format_answer_text(doc, answer)

        # Divider line
        doc.add_paragraph("--------------------------------------------------\n\n")

    doc.save(output_file)
    return output_file_name


def format_answer_text(paragraph, answer_text):
    paragraphs = answer_text.split("\n\n")
    
    for para in paragraphs:
        lines = para.splitlines()
        
        for line in lines:
            # Create a new paragraph for each line
            current_paragraph = paragraph.add_paragraph()
            
            # Handle headers starting with ###
            if line.startswith("### "):
                header_run = current_paragraph.add_run(line[4:])
                header_run.bold = True
                header_run.font.size = Pt(14)
                current_paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            # Primary bullet points
            elif line.startswith("1. ") or line.startswith("- "):
                current_paragraph.paragraph_format.left_indent = Pt(20)
                process_text_with_bold(current_paragraph, line)

            # Secondary bullet points with extra indentation
            elif line.startswith("  - "):
                current_paragraph.paragraph_format.left_indent = Pt(40)
                process_text_with_bold(current_paragraph, line)
            
            # Regular text
            else:
                process_text_with_bold(current_paragraph, line)


def process_text_with_bold(paragraph, text):
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            bold_run = paragraph.add_run(part[2:-2])  # Remove the ** markers
            bold_run.bold = True
            bold_run.font.size = Pt(11)
        else:
            normal_run = paragraph.add_run(part)
            normal_run.font.size = Pt(11)



@app.get("/task_status/{task_id}")
async def task_status(task_id: str):
    task = celery.AsyncResult(task_id)
    if task.state == "PENDING":
        return {"status": "pending"}
    elif task.state == "FAILURE":
        return {"status": "failed", "error": str(task.info)}
    elif task.state == "SUCCESS":
        result = task.result
        return {
            "status": result["status"],
            "output_file": result.get("output_file"),
        }
    else:
        return {"status": "unknown"}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app:app", host="0.0.0.0", port=8080, reload=True)